"""
Local Self-Hosted PageIndex Service

Integrates VectifyAI's PageIndex for vectorless, reasoning-based RAG.
Builds hierarchical tree indexes from PDFs/text and searches them
using LLM reasoning — no vector database needed.

Usage:
    service = LocalPageIndexService()
    # Index a PDF
    result = await service.index_document("/path/to/report.pdf", {"deal_id": "123"})
    # Query across all indexes
    chunks = await service.query("What is the company's revenue growth?")
"""

import os
import json
import uuid
import asyncio
import shutil
from pathlib import Path
from typing import Dict, Any, List, Optional
from datetime import datetime
from dataclasses import dataclass, field, asdict
import structlog

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None

logger = structlog.get_logger()


# ===== Data Models =====


@dataclass
class IndexedDocument:
    """Metadata for an indexed document"""

    doc_id: str
    index_id: str
    filename: str
    file_path: str
    file_type: str
    tree_path: str  # path to the generated tree JSON
    total_pages: int
    total_nodes: int
    metadata: Dict[str, Any]
    created_at: str
    status: str = "indexed"  # indexed, error, processing


@dataclass
class TreeNode:
    """A node in the PageIndex tree"""

    node_id: str
    title: str
    summary: str
    start_page: int
    end_page: int
    content: str = ""
    children: List["TreeNode"] = field(default_factory=list)


@dataclass
class SearchResult:
    """A search result from the local index"""

    chunk_id: str
    content: str
    page_number: int
    node_title: str
    relevance_score: float
    doc_id: str
    metadata: Dict[str, Any] = field(default_factory=dict)


# ===== Local PageIndex Service =====


class LocalPageIndexService:
    """
    Self-hosted PageIndex service for DealForge AI.

    Manages document indexing, tree storage, and reasoning-based retrieval
    without requiring the VectifyAI cloud API.
    """

    def __init__(
        self,
        storage_dir: str = None,
        model: str = "gpt-4o",
        openai_api_key: str = None,
        gemini_api_key: str = None,
        use_gemini: bool = False,
    ):
        self.storage_dir = Path(storage_dir or self._default_storage_dir())
        self.indexes_dir = self.storage_dir / "indexes"
        self.trees_dir = self.storage_dir / "trees"
        self.docs_dir = self.storage_dir / "documents"
        self.catalog_path = self.storage_dir / "catalog.json"

        self.model = model
        self.openai_api_key = openai_api_key
        self.gemini_api_key = gemini_api_key
        self.use_gemini = use_gemini

        # Ensure directories exist
        for d in [self.storage_dir, self.indexes_dir, self.trees_dir, self.docs_dir]:
            d.mkdir(parents=True, exist_ok=True)

        # Load catalog
        self.catalog: Dict[str, IndexedDocument] = self._load_catalog()

        logger.info(
            "LocalPageIndexService initialized",
            storage_dir=str(self.storage_dir),
            indexed_docs=len(self.catalog),
        )

    @staticmethod
    def _default_storage_dir() -> str:
        """Default storage location"""
        return os.path.join(
            os.path.dirname(os.path.abspath(__file__)),
            "..",
            "..",
            "..",
            "pageindex_data",
        )

    # ===== Catalog Management =====

    def _load_catalog(self) -> Dict[str, IndexedDocument]:
        """Load the document catalog from disk"""
        if self.catalog_path.exists():
            try:
                with open(self.catalog_path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                catalog = {}
                for doc_id, doc_data in data.items():
                    catalog[doc_id] = IndexedDocument(**doc_data)
                return catalog
            except Exception as e:
                logger.error("Failed to load catalog", error=str(e))
        return {}

    def _save_catalog(self):
        """Persist catalog to disk"""
        data = {doc_id: asdict(doc) for doc_id, doc in self.catalog.items()}
        with open(self.catalog_path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)

    # ===== Document Indexing =====

    async def index_document(
        self,
        file_path: str,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> IndexedDocument:
        """
        Index a document using PageIndex locally.

        Supports: PDF, MD, TXT, DOCX (txt/docx converted to md first)

        Args:
            file_path: Path to the document
            metadata: Optional metadata (deal_id, tags, etc.)

        Returns:
            IndexedDocument with tree details
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")

        doc_id = str(uuid.uuid4())
        index_id = f"idx_{doc_id[:8]}"
        file_type = file_path.suffix.lower()

        logger.info(
            "Indexing document", doc_id=doc_id, file=str(file_path), type=file_type
        )

        # Copy file to our storage
        stored_path = self.docs_dir / f"{doc_id}{file_type}"
        shutil.copy2(str(file_path), str(stored_path))

        # Build the tree index
        tree_data = await self._build_tree(str(stored_path), file_type)

        # Save tree
        tree_path = self.trees_dir / f"{doc_id}_tree.json"
        with open(tree_path, "w", encoding="utf-8") as f:
            json.dump(tree_data, f, indent=2, ensure_ascii=False)

        # Count nodes
        total_nodes = self._count_nodes(tree_data)
        total_pages = self._get_max_page(tree_data)

        # Create catalog entry
        doc = IndexedDocument(
            doc_id=doc_id,
            index_id=index_id,
            filename=file_path.name,
            file_path=str(stored_path),
            file_type=file_type,
            tree_path=str(tree_path),
            total_pages=total_pages,
            total_nodes=total_nodes,
            metadata=metadata or {},
            created_at=datetime.utcnow().isoformat(),
            status="indexed",
        )

        self.catalog[doc_id] = doc
        self._save_catalog()

        logger.info(
            "Document indexed successfully",
            doc_id=doc_id,
            nodes=total_nodes,
            pages=total_pages,
        )

        return doc

    async def index_text(
        self,
        content: str,
        title: str = "Untitled",
        metadata: Optional[Dict[str, Any]] = None,
    ) -> IndexedDocument:
        """
        Index raw text content (e.g., scraped web content, notes).

        Saves as markdown and indexes using PageIndex tree builder.
        """
        doc_id = str(uuid.uuid4())
        index_id = f"idx_{doc_id[:8]}"

        # Save text as markdown
        md_path = self.docs_dir / f"{doc_id}.md"
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(f"# {title}\n\n{content}")

        # Build tree from markdown
        tree_data = await self._build_tree(str(md_path), ".md")

        # Save tree
        tree_path = self.trees_dir / f"{doc_id}_tree.json"
        with open(tree_path, "w", encoding="utf-8") as f:
            json.dump(tree_data, f, indent=2, ensure_ascii=False)

        total_nodes = self._count_nodes(tree_data)

        doc = IndexedDocument(
            doc_id=doc_id,
            index_id=index_id,
            filename=f"{title}.md",
            file_path=str(md_path),
            file_type=".md",
            tree_path=str(tree_path),
            total_pages=1,
            total_nodes=total_nodes,
            metadata=metadata or {},
            created_at=datetime.utcnow().isoformat(),
            status="indexed",
        )

        self.catalog[doc_id] = doc
        self._save_catalog()

        logger.info("Text indexed", doc_id=doc_id, title=title, nodes=total_nodes)
        return doc

    async def _build_tree(self, file_path: str, file_type: str) -> Dict:
        """Build PageIndex tree from a document"""
        try:
            # Try using the pageindex package
            from pageindex import config as pi_config, page_index_main
            from pageindex.page_index_md import md_to_tree

            if file_type == ".pdf":
                # Set OpenAI API key for PageIndex
                if self.openai_api_key:
                    os.environ["CHATGPT_API_KEY"] = self.openai_api_key

                opt = pi_config(
                    model=self.model,
                    toc_check_page_num=20,
                    max_page_num_each_node=10,
                    max_token_num_each_node=20000,
                    if_add_node_id="yes",
                    if_add_node_summary="yes",
                    if_add_doc_description="yes",
                    if_add_node_text="yes",
                )

                # Run in an executor since page_index_main may be blocking
                loop = asyncio.get_event_loop()
                tree_data = await loop.run_in_executor(
                    None, page_index_main, file_path, opt
                )
                return tree_data

            elif file_type in (".md", ".markdown"):
                tree_data = await md_to_tree(
                    md_path=file_path,
                    if_thinning=False,
                    if_add_node_summary="yes",
                    model=self.model,
                    if_add_doc_description="yes",
                    if_add_node_text="yes",
                    if_add_node_id="yes",
                )
                return tree_data

            elif file_type == ".txt":
                # Convert to markdown first
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
                md_path = file_path.replace(".txt", ".md")
                with open(md_path, "w", encoding="utf-8") as f:
                    f.write(f"# Document\n\n{text}")
                return await self._build_tree(md_path, ".md")

            else:
                # Fallback: create a simple tree manually
                return self._build_simple_tree(file_path)

        except ImportError:
            logger.warning(
                "PageIndex package not installed, using fallback tree builder"
            )
            return self._build_simple_tree(file_path)

    def _build_simple_tree(self, file_path: str) -> Dict:
        """
        Fallback tree builder when PageIndex package is not available.
        Creates a basic hierarchical structure from file content.
        """
        content = ""
        file_ext = os.path.splitext(file_path)[1].lower()

        try:
            if file_ext == ".pdf":
                if pypdf:
                    try:
                        with open(file_path, "rb") as f:
                            reader = pypdf.PdfReader(f)
                            text_parts = []
                            for page in reader.pages:
                                part = page.extract_text()
                                if part:
                                    text_parts.append(part)
                            content = "\n\n".join(text_parts)
                    except Exception as e:
                        logger.error("Error extracting text from PDF", error=str(e))
                        content = f"[Error reading PDF: {os.path.basename(file_path)}]"
                else:
                    content = f"[Binary file (pypdf not installed): {os.path.basename(file_path)}]"

            elif file_ext == ".docx":
                if docx:
                    try:
                        doc = docx.Document(file_path)
                        content = "\n".join([para.text for para in doc.paragraphs])
                    except Exception as e:
                        logger.error("Error extracting text from DOCX", error=str(e))
                        content = f"[Error reading DOCX: {os.path.basename(file_path)}]"
                else:
                    content = f"[Binary file (python-docx not installed): {os.path.basename(file_path)}]"

            else:
                # Assume text/markdown
                try:
                    with open(file_path, "r", encoding="utf-8") as f:
                        content = f.read()
                except UnicodeDecodeError:
                    content = f"[Binary file: {os.path.basename(file_path)}]"

        except Exception as e:
            logger.error("Fallback extraction failed", error=str(e))
            content = f"[Error processing file: {os.path.basename(file_path)}]"

        if not content.strip():
            content = f"[Empty or unreadable file: {os.path.basename(file_path)}]"

        # Split into chunks of ~2000 chars
        chunks = []
        words = content.split()
        current_chunk = []
        current_len = 0

        for word in words:
            current_chunk.append(word)
            current_len += len(word) + 1
            if current_len >= 2000:
                chunks.append(" ".join(current_chunk))
                current_chunk = []
                current_len = 0

        if current_chunk:
            chunks.append(" ".join(current_chunk))

        # Build a flat tree
        nodes = []
        for i, chunk in enumerate(chunks):
            nodes.append(
                {
                    "node_id": f"{i:04d}",
                    "title": f"Section {i + 1}",
                    "summary": chunk[:200] + "..." if len(chunk) > 200 else chunk,
                    "start_index": i,
                    "end_index": i + 1,
                    "text": chunk,
                    "nodes": [],
                }
            )

        return {
            "title": os.path.basename(file_path),
            "description": f"Indexed from {os.path.basename(file_path)}",
            "total_pages": len(chunks),
            "nodes": nodes,
        }

    # ===== Querying =====

    async def query(
        self,
        query: str,
        index_id: Optional[str] = None,
        deal_id: Optional[str] = None,
        top_k: int = 5,
    ) -> List[SearchResult]:
        """
        Search across indexed documents using reasoning-based retrieval.

        Walks the tree structure to find relevant sections.

        Args:
            query: Search query
            index_id: Specific index to search (optional)
            deal_id: Filter by deal (optional)
            top_k: Number of results to return

        Returns:
            List of SearchResult objects
        """
        results: List[SearchResult] = []

        # Determine which documents to search
        docs_to_search = []
        for doc_id, doc in self.catalog.items():
            if doc.status != "indexed":
                continue
            if index_id and doc.index_id != index_id:
                continue
            if deal_id and doc.metadata.get("deal_id") != deal_id:
                continue
            docs_to_search.append(doc)

        if not docs_to_search:
            logger.info("No documents to search", query=query)
            return []

        # Search each document's tree
        for doc in docs_to_search:
            try:
                tree_path = Path(doc.tree_path)
                if not tree_path.exists():
                    continue

                with open(tree_path, "r", encoding="utf-8") as f:
                    tree_data = json.load(f)

                # Extract matching nodes from tree
                doc_results = self._search_tree(query, tree_data, doc)
                results.extend(doc_results)

            except Exception as e:
                logger.error(
                    "Error searching document",
                    doc_id=doc.doc_id,
                    error=str(e),
                )

        # Sort by relevance and limit
        results.sort(key=lambda r: r.relevance_score, reverse=True)
        return results[:top_k]

    def _search_tree(
        self,
        query: str,
        tree_data: Dict,
        doc: IndexedDocument,
    ) -> List[SearchResult]:
        """
        Search a tree structure for relevant nodes using keyword matching.

        Note: For full reasoning-based search, the PageIndex package's
        tree search with LLM would be used. This is a fast local fallback
        that uses keyword overlap scoring.
        """
        query_words = set(query.lower().split())
        results = []

        def walk_nodes(nodes, depth=0):
            if not isinstance(nodes, list):
                return
            for node in nodes:
                if not isinstance(node, dict):
                    continue

                # Get searchable text
                title = node.get("title", "")
                summary = node.get("summary", "")
                text = node.get("text", "")
                searchable = f"{title} {summary} {text}".lower()

                # Calculate relevance
                searchable_words = set(searchable.split())
                overlap = query_words & searchable_words
                if overlap:
                    relevance = len(overlap) / max(len(query_words), 1)
                    # Boost exact phrase matches
                    if query.lower() in searchable:
                        relevance = min(1.0, relevance + 0.3)
                    # Boost title matches
                    if any(w in title.lower() for w in query_words):
                        relevance = min(1.0, relevance + 0.2)

                    content = text if text else summary

                    if content:
                        results.append(
                            SearchResult(
                                chunk_id=f"{doc.doc_id}_{node.get('node_id', 'unknown')}",
                                content=content[:2000],  # Cap content length
                                page_number=node.get("start_index", 0),
                                node_title=title,
                                relevance_score=round(relevance, 4),
                                doc_id=doc.doc_id,
                                metadata={
                                    "filename": doc.filename,
                                    "depth": depth,
                                    **doc.metadata,
                                },
                            )
                        )

                # Recurse into children
                walk_nodes(node.get("nodes", []), depth + 1)

        walk_nodes(tree_data.get("nodes", []))
        return results

    # ===== Management =====

    def list_documents(self, deal_id: Optional[str] = None) -> List[IndexedDocument]:
        """List all indexed documents, optionally filtered by deal"""
        docs = list(self.catalog.values())
        if deal_id:
            docs = [d for d in docs if d.metadata.get("deal_id") == deal_id]
        return docs

    def get_document(self, doc_id: str) -> Optional[IndexedDocument]:
        """Get a specific document by ID"""
        return self.catalog.get(doc_id)

    def delete_document(self, doc_id: str) -> bool:
        """Delete a document and its index"""
        doc = self.catalog.get(doc_id)
        if not doc:
            return False

        # Remove files
        for path_str in [doc.file_path, doc.tree_path]:
            try:
                path = Path(path_str)
                if path.exists():
                    path.unlink()
            except Exception as e:
                logger.error("Failed to delete file", path=path_str, error=str(e))

        # Remove from catalog
        del self.catalog[doc_id]
        self._save_catalog()

        logger.info("Document deleted", doc_id=doc_id)
        return True

    def get_stats(self) -> Dict[str, Any]:
        """Get storage statistics"""
        total_docs = len(self.catalog)
        total_nodes = sum(d.total_nodes for d in self.catalog.values())
        total_pages = sum(d.total_pages for d in self.catalog.values())

        # Storage size
        total_size = 0
        for d in self.catalog.values():
            for p in [d.file_path, d.tree_path]:
                try:
                    total_size += Path(p).stat().st_size
                except OSError:
                    pass

        return {
            "total_documents": total_docs,
            "total_nodes": total_nodes,
            "total_pages": total_pages,
            "storage_bytes": total_size,
            "storage_mb": round(total_size / (1024 * 1024), 2),
            "storage_dir": str(self.storage_dir),
        }

    # ===== Utility =====

    def _count_nodes(self, tree_data: Dict) -> int:
        """Count total nodes in a tree"""
        count = 0

        def walk(nodes):
            nonlocal count
            if not isinstance(nodes, list):
                return
            for node in nodes:
                if isinstance(node, dict):
                    count += 1
                    walk(node.get("nodes", []))

        walk(tree_data.get("nodes", []))
        return count

    def _get_max_page(self, tree_data: Dict) -> int:
        """Get the maximum page number from a tree"""
        max_page = 0

        def walk(nodes):
            nonlocal max_page
            if not isinstance(nodes, list):
                return
            for node in nodes:
                if isinstance(node, dict):
                    end = node.get("end_index", 0)
                    if isinstance(end, int) and end > max_page:
                        max_page = end
                    walk(node.get("nodes", []))

        walk(tree_data.get("nodes", []))
        return max_page


# ===== Singleton =====

_local_pageindex: Optional[LocalPageIndexService] = None


def get_local_pageindex(
    storage_dir: str = None,
    **kwargs,
) -> LocalPageIndexService:
    """Get or create the local PageIndex service singleton"""
    global _local_pageindex
    if _local_pageindex is None:
        _local_pageindex = LocalPageIndexService(storage_dir=storage_dir, **kwargs)
    return _local_pageindex
